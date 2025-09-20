# utils/exporter.py
import re
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import html

def _parse_and_add_line_docx(paragraph, line):
    """
    Parses a line of markdown for bold, italic, and footnotes, 
    and adds formatted runs to a docx paragraph.
    """
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\[\^\d+\])', line)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif re.match(r'\[\^(\d+)\]', part):
            note_num = re.match(r'\[\^(\d+)\]', part).group(1)
            run = paragraph.add_run(note_num)
            run.font.superscript = True
        elif part:
            paragraph.add_run(part)

def export_to_docx(markdown_text, footnotes, filename):
    """
    Exports a markdown string to a .docx file with endnotes.
    - footnotes: A dict where keys are footnote numbers (str) and values are the note text (str).
    """
    document = Document()
    
    section = document.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    for line in markdown_text.split('\n'):
        stripped_line = line.strip()
        if stripped_line.startswith('### '):
            p = document.add_heading(level=3)
            _parse_and_add_line_docx(p, stripped_line[4:])
        elif stripped_line.startswith('## '):
            p = document.add_heading(level=2)
            _parse_and_add_line_docx(p, stripped_line[3:])
        elif stripped_line.startswith('# '):
            p = document.add_heading(level=1)
            _parse_and_add_line_docx(p, stripped_line[2:])
        elif stripped_line:
            p = document.add_paragraph()
            _parse_and_add_line_docx(p, stripped_line)
        else:
            document.add_paragraph()
            
    if footnotes:
        document.add_page_break()
        document.add_heading('Notes', level=1)
        sorted_notes = sorted(footnotes.items(), key=lambda item: int(item[0]))
        for num, text in sorted_notes:
            p = document.add_paragraph()
            p.add_run(f"{num}. ").bold = True
            for line in text.split('\n'):
                p.add_run(line).add_break()

    document.save(filename)

def export_to_pdf(markdown_text, footnotes, filename):
    """
    Exports a markdown string to a .pdf file using an HTML-based conversion.
    - footnotes: A dict where keys are footnote numbers (str) and values are the note text (str).
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(left=12.7, top=12.7, right=12.7)
    
    # --- MODIFICATION: Add fallback for missing fonts ---
    try:
        pdf.add_font("DejaVu", "", "DejaVuSans.ttf")
        pdf.add_font("DejaVu", "B", "DejaVuSans-Bold.ttf")
        pdf.add_font("DejaVu", "I", "DejaVuSans-Oblique.ttf")
        pdf.add_font("DejaVu", "BI", "DejaVuSans-BoldOblique.ttf")
        pdf.set_font("DejaVu", size=12)
    except FileNotFoundError:
        print("WARNING: DejaVu font not found. Falling back to standard font. Unicode characters may not render correctly.")
        pdf.set_font("Helvetica", size=12)
    # --- END MODIFICATION ---

    html_content = ""
    for line in markdown_text.split('\n'):
        line = html.escape(line)
        line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
        line = re.sub(r'\*(.*?)\*', r'<i>\1</i>', line)
        line = re.sub(r'\[\^(\d+)\]', r'<sup>\1</sup>', line)

        if line.strip().startswith('# '):
            html_content += f"<h1>{line.strip()[2:]}</h1>"
        elif line.strip().startswith('## '):
            html_content += f"<h2>{line.strip()[3:]}</h2>"
        elif line.strip().startswith('### '):
            html_content += f"<h3>{line.strip()[4:]}</h3>"
        else:
            html_content += f"<p>{line}</p>"
            
    if footnotes:
        html_content += "<hr>"
        html_content += "<h2>Notes</h2>"
        sorted_notes = sorted(footnotes.items(), key=lambda item: int(item[0]))
        for num, text in sorted_notes:
            text_html = html.escape(text).replace('\n', '<br>')
            html_content += f"<p><b>{num}.</b> {text_html}</p>"

    pdf.write_html(html_content)
    pdf.output(filename)
